"""python-docx renderers for the tailored resume and cover letter.

Layout is ATS-safe: Calibri 10/11/18pt, no tables, no colors, no icons, one
page-sized section. We write to a sibling ``.tmp`` path first and rename at
the end so a crash can never leave a half-written .docx behind.
"""
from __future__ import annotations

import os
from datetime import date
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree

from src.models.job import Job
from src.writing.resume_tailor import SelectedHighlight, TailoredResume

_FONT = "Calibri"

_FR_MONTHS = [
    "janvier", "février", "mars", "avril", "mai", "juin",
    "juillet", "août", "septembre", "octobre", "novembre", "décembre",
]
_EN_MONTHS = [
    "January", "February", "March", "April", "May", "June",
    "July", "August", "September", "October", "November", "December",
]


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
def _configure_page(doc: _Document) -> None:
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = etree.SubElement(rpr, qn("w:rFonts"))
    for attr in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(attr), _FONT)


def _run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
         size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = _FONT
    if size is not None:
        run.font.size = Pt(size)


def _section_heading(doc: _Document, label: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    _run(p, label.upper(), bold=True, size=11)

    pPr = p._p.get_or_add_pPr()
    existing = pPr.find(qn("w:pBdr"))
    if existing is not None:
        pPr.remove(existing)
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "808080")


def _bullet(doc: _Document, text: str) -> None:
    if not text:
        return
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    _run(p, text, size=10)


def _atomic_save(doc: _Document, output_path: Path) -> Path:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_suffix(output_path.suffix + ".tmp")
    doc.save(str(tmp_path))
    os.replace(tmp_path, output_path)
    return output_path


def _format_date(language: str, when: date | None = None) -> str:
    today = when or date.today()
    months = _FR_MONTHS if language == "fr" else _EN_MONTHS
    if language == "fr":
        return f"{today.day} {months[today.month - 1]} {today.year}"
    return f"{months[today.month - 1]} {today.day}, {today.year}"


# ---------------------------------------------------------------------------
# resume
# ---------------------------------------------------------------------------
def _render_header(doc: _Document, profile: dict) -> None:
    personal: dict[str, Any] = profile.get("personal", {})
    name = personal.get("name", "")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    _run(p, name, bold=True, size=18)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    _run(p2, "Data Analyst | Analytics & Automation", size=12)

    contact_parts = [
        personal.get("name", ""),
        personal.get("location", ""),
        personal.get("phone", ""),
        personal.get("email", ""),
        personal.get("linkedin", ""),
        personal.get("github", ""),
        personal.get("website", ""),
    ]
    contact_line = " · ".join([c for c in contact_parts if c])
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(4)
    _run(p3, contact_line, size=10)

    auth_stmt = profile.get("work_authorization", {}).get("resume_statement", "")
    if auth_stmt:
        p4 = doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(6)
        _run(p4, auth_stmt, italic=True, size=9)


def _render_summary(doc: _Document, tailored: TailoredResume) -> None:
    _section_heading(doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _run(p, tailored.summary, size=10)


def _render_technical_skills(doc: _Document, profile: dict) -> None:
    _section_heading(doc, "Technical Skills")
    skills = profile.get("skills", {})
    languages = profile.get("languages", [])

    rows: list[tuple[str, list[str]]] = [
        ("Languages & Query", skills.get("languages_code", [])),
        ("Data & BI", skills.get("data_bi", [])),
        ("Databases", skills.get("databases", [])),
        ("Automation", skills.get("automation", [])),
        ("Cloud & Platforms", skills.get("cloud_platforms", [])),
        ("AI Tooling", skills.get("ai_tooling", [])),
        ("Languages Spoken", languages),
    ]
    for label, values in rows:
        if not values:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, f"{label}: ", bold=True, size=10)
        _run(p, ", ".join(values), size=10)


def _date_range_for(profile: dict, company: str, title: str) -> str:
    for exp in profile.get("experience", []):
        if exp.get("company") == company and exp.get("title") == title:
            start = exp.get("start", "")
            end = exp.get("end", "")
            if start and end:
                return f"{start} – {end}"
            return start or end or ""
    return ""


def _location_for(profile: dict, company: str, title: str) -> str:
    for exp in profile.get("experience", []):
        if exp.get("company") == company and exp.get("title") == title:
            return exp.get("location", "")
    return ""


def _render_experience(
    doc: _Document,
    profile: dict,
    tailored: TailoredResume,
) -> None:
    _section_heading(doc, "Professional Experience")
    if not tailored.selected_highlights:
        return

    # preserve order of first appearance per (company, title)
    seen: list[tuple[str, str]] = []
    by_key: dict[tuple[str, str], list[SelectedHighlight]] = {}
    for hl in tailored.selected_highlights:
        key = (hl.company, hl.title)
        if key not in by_key:
            by_key[key] = []
            seen.append(key)
        by_key[key].append(hl)

    for (company, title) in seen:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(
            Inches(7.1),
            alignment=WD_TAB_ALIGNMENT.RIGHT,
            leader=WD_TAB_LEADER.SPACES,
        )
        location = _location_for(profile, company, title)
        header_text = f"{company}" + (f" — {location}" if location else "")
        date_text = _date_range_for(profile, company, title)
        _run(p, header_text, bold=True, size=10)
        _run(p, "\t" + date_text, size=10)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, title, italic=True, size=10)

        for hl in by_key[(company, title)]:
            _bullet(doc, hl.text)


def _render_projects(
    doc: _Document,
    profile: dict,
    tailored: TailoredResume,
) -> None:
    _section_heading(doc, "Selected Projects")
    by_name = {p["name"]: p for p in profile.get("projects", [])}
    for name in tailored.selected_projects:
        project = by_name.get(name)
        if not project:
            continue
        text = f"{name} — {project.get('description', '')}"
        _bullet(doc, text)


def _render_education(doc: _Document, profile: dict) -> None:
    _section_heading(doc, "Education")
    for entry in profile.get("education", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        _run(p, entry.get("institution", ""), bold=True, size=10)
        detail = (
            f" — {entry.get('credential', '')}"
            f" ({entry.get('year', '')})"
        )
        _run(p, detail, size=10)
        loc = entry.get("location")
        if loc:
            _run(p, f" · {loc}", italic=True, size=10)


def _render_certifications(doc: _Document, profile: dict) -> None:
    certs = profile.get("certifications", [])
    if not certs:
        return
    _section_heading(doc, "Certifications")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    _run(p, " · ".join(str(c) for c in certs), size=10)


def build_tailored_resume_docx(
    profile: dict,
    tailored: TailoredResume,
    output_path: Path,
) -> Path:
    """Render the tailored resume to ``output_path`` and return the path."""
    doc = Document()
    _configure_page(doc)
    _render_header(doc, profile)
    _render_summary(doc, tailored)
    _render_technical_skills(doc, profile)
    _render_experience(doc, profile, tailored)
    _render_projects(doc, profile, tailored)
    _render_education(doc, profile)
    _render_certifications(doc, profile)
    return _atomic_save(doc, output_path)


# ---------------------------------------------------------------------------
# cover letter
# ---------------------------------------------------------------------------
def _cover_greeting(language: str, company: str) -> str:
    if language == "fr":
        return "Madame, Monsieur,"
    return f"Dear Hiring Team at {company},"


def _cover_closing(language: str) -> str:
    if language == "fr":
        return "Cordialement,"
    return "Sincerely,"


def _split_paragraphs(text: str) -> list[str]:
    blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    if len(blocks) >= 2:
        return blocks
    # LLM likely used single newlines — fall back to line splits.
    return [line.strip() for line in text.split("\n") if line.strip()]


def build_cover_letter_docx(
    profile: dict,
    job: Job,
    cover_text: str,
    output_path: Path,
    language: str = "en",
) -> Path:
    """Render the cover letter. ``cover_text`` is the raw body produced by the
    LLM (greeting / closing will be re-applied here for layout consistency)."""
    doc = Document()
    _configure_page(doc)
    personal = profile.get("personal", {})

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(6)
    _run(date_p, _format_date(language), size=10)

    name_p = doc.add_paragraph()
    name_p.paragraph_format.space_after = Pt(0)
    _run(name_p, personal.get("name", ""), bold=True, size=11)
    contact_parts = [
        personal.get("location", ""),
        personal.get("phone", ""),
        personal.get("email", ""),
        personal.get("linkedin", ""),
    ]
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(10)
    _run(contact_p, " · ".join(c for c in contact_parts if c), size=9)

    company_p = doc.add_paragraph()
    company_p.paragraph_format.space_after = Pt(10)
    _run(company_p, job.company, bold=True, size=10)

    greeting_p = doc.add_paragraph()
    greeting_p.paragraph_format.space_after = Pt(8)
    _run(greeting_p, _cover_greeting(language, job.company), size=10)

    # Strip any greeting / closing the LLM already emitted so we don't duplicate.
    body = cover_text or ""
    body_lines = body.splitlines()
    trimmed: list[str] = []
    banned_leads = ("dear hiring", "madame,", "madame, monsieur")
    banned_tails = ("sincerely,", "cordialement,", "davi almeida")
    for line in body_lines:
        low = line.strip().lower()
        if not trimmed and any(low.startswith(p) for p in banned_leads):
            continue
        if any(low == t or low.startswith(t) for t in banned_tails):
            continue
        trimmed.append(line)
    cleaned_body = "\n".join(trimmed).strip() or body.strip()

    for block in _split_paragraphs(cleaned_body):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        _run(p, block, size=10)

    close_p = doc.add_paragraph()
    close_p.paragraph_format.space_after = Pt(0)
    _run(close_p, _cover_closing(language), size=10)
    name_close_p = doc.add_paragraph()
    _run(name_close_p, personal.get("name", ""), size=10)

    return _atomic_save(doc, output_path)
