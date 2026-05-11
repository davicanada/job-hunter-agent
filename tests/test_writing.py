"""Writing pipeline unit tests — pure logic, no LLM or network."""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from src.models.job import Job
from src.writing.cover_letter import detect_language
from src.writing.docx_builder import (
    build_cover_letter_docx,
    build_tailored_resume_docx,
)
from src.writing.paths import make_job_output_paths, slugify
from src.writing.resume_tailor import SelectedHighlight, TailoredResume


def _job(**overrides) -> Job:
    base = {
        "external_id": "test-writing",
        "source": "remoteok",
        "title": "Data Analyst",
        "company": "Shopify Inc.",
        "description": "Looking for a data analyst with SQL and Python.",
        "url": "https://example.com/job/123",
        "allows_target_region": True,
    }
    base.update(overrides)
    return Job(**base)


def _tiny_profile() -> dict:
    return {
        "personal": {
            "name": "Davi Almeida",
            "location": "Montreal, QC",
            "email": "davi@example.com",
            "phone": "555-1234",
            "linkedin": "linkedin.com/in/davi",
            "github": "github.com/davi",
            "website": "davi.dev",
        },
        "work_authorization": {
            "resume_statement": "Authorized to work in Canada — PR approved, COPR pending.",
        },
        "languages": ["English", "French"],
        "skills": {
            "languages_code": ["SQL", "Python"],
            "data_bi": ["Power BI"],
            "databases": ["PostgreSQL"],
            "automation": ["Playwright"],
            "cloud_platforms": ["Supabase"],
            "ai_tooling": ["Groq"],
        },
        "experience": [
            {
                "company": "Otodata",
                "title": "Documentation Technician",
                "start": "2024",
                "end": "present",
                "location": "Montreal, QC",
                "highlights": [
                    {"summary": "Automated monthly AVL report in VBA.", "tags": ["vba"]},
                ],
            }
        ],
        "projects": [
            {"name": "Taskor", "description": "Full-stack performance app.", "stack": ["Next.js"]},
        ],
        "education": [
            {"institution": "Trebas Institute", "credential": "ACS Analytics", "year": "2023", "location": "Montreal, QC"}
        ],
        "certifications": ["Google Data Analytics (2023)"],
    }


# ---------------------------------------------------------------------------
# slugify / paths
# ---------------------------------------------------------------------------
def test_slugify_company():
    assert slugify("Shopify Inc.") == "shopify_inc"


def test_slugify_strips_unicode_and_punct():
    assert slugify("Données & Analytics!") == "donn_es_analytics"


def test_slugify_empty_is_untitled():
    assert slugify("") == "untitled"


def test_slugify_trims_to_max_len():
    long = "a" * 100
    out = slugify(long, max_len=40)
    assert len(out) <= 40


def test_make_paths_stable_for_same_job(tmp_path: Path):
    job = _job(company="Acme Corp", title="Data Analyst")
    resume_a, cover_a = make_job_output_paths(tmp_path, job)
    resume_b, cover_b = make_job_output_paths(tmp_path, job)
    assert resume_a == resume_b
    assert cover_a == cover_b
    assert resume_a.name.endswith("_resume.docx")
    assert cover_a.name.endswith("_cover.docx")


# ---------------------------------------------------------------------------
# language detection
# ---------------------------------------------------------------------------
def test_language_detection_french():
    desc = (
        "Nous recherchons un développeur passionné pour rejoindre notre "
        "équipe. Vous travaillerez avec des données chez une entreprise "
        "dynamique."
    )
    assert detect_language(desc) == "fr"


def test_language_detection_english():
    desc = (
        "We are looking for a passionate data analyst to join our team. "
        "You will work with SQL and Python and build dashboards in Power BI."
    )
    assert detect_language(desc) == "en"


def test_language_detection_empty_defaults_to_english():
    assert detect_language("") == "en"
    assert detect_language(None) == "en"


# ---------------------------------------------------------------------------
# docx builder smoke
# ---------------------------------------------------------------------------
def test_docx_builder_smoke(tmp_path: Path):
    profile = _tiny_profile()
    tailored = TailoredResume(
        summary="Data analyst focused on SQL, Python, and BI automation.",
        selected_highlights=[
            SelectedHighlight(
                company="Otodata",
                title="Documentation Technician",
                text="Automated monthly AVL report in VBA, cutting a 1h task to <1min.",
            )
        ],
        selected_projects=["Taskor"],
        keywords_added=["SQL", "Python"],
    )
    resume_path = tmp_path / "resume.docx"
    cover_path = tmp_path / "cover.docx"

    out_resume = build_tailored_resume_docx(profile, tailored, resume_path)
    assert out_resume.exists()
    assert out_resume.stat().st_size > 2000  # real docx payload
    Document(str(out_resume))  # python-docx can re-open

    out_cover = build_cover_letter_docx(
        profile,
        _job(),
        cover_text=(
            "Dear Hiring Team at Shopify Inc.,\n\n"
            "Opening paragraph with concrete hook about Shopify and Davi's work.\n\n"
            "Evidence paragraph citing Otodata automation metrics.\n\n"
            "Closing paragraph with genuine enthusiasm and the authorization line.\n\n"
            "Sincerely,\nDavi Almeida"
        ),
        output_path=cover_path,
        language="en",
    )
    assert out_cover.exists()
    assert out_cover.stat().st_size > 2000
    Document(str(out_cover))


def test_docx_builder_french_cover(tmp_path: Path):
    profile = _tiny_profile()
    path = tmp_path / "cover_fr.docx"
    out = build_cover_letter_docx(
        profile,
        _job(company="Acme FR"),
        cover_text=(
            "Madame, Monsieur,\n\n"
            "Premier paragraphe sur l'entreprise et Davi.\n\n"
            "Deuxième paragraphe avec des exemples concrets.\n\n"
            "Troisième paragraphe de clôture.\n\n"
            "Cordialement,\nDavi Almeida"
        ),
        output_path=path,
        language="fr",
    )
    assert out.exists()
    Document(str(out))


def test_profile_fixture_is_json_serialisable():
    # defensive: ensure our test profile stays in-sync with what the docx
    # builder expects when the real profile.json evolves.
    assert json.dumps(_tiny_profile())
